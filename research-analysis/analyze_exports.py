from __future__ import annotations

import argparse
import json
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from scipy import stats
from statsmodels.formula.api import ols
from statsmodels.stats.anova import anova_lm

try:
    from docx import Document
except ImportError:  # pragma: no cover - optional dependency during bootstrap
    Document = None


PARTICIPANT_NUMERIC_COLUMNS = [
    "analysisCount",
    "weeklySessionCount",
    "pretestPitch",
    "posttestPitch",
    "pretestRhythm",
    "posttestRhythm",
    "pitchGain",
    "rhythmGain",
    "usefulness",
    "easeOfUse",
    "feedbackClarity",
    "confidence",
    "continuance",
    "questionnaireCount",
    "expertPretestPitch",
    "expertPosttestPitch",
    "expertPretestRhythm",
    "expertPosttestRhythm",
]

QUESTIONNAIRE_NUMERIC_COLUMNS = ["usefulness", "easeOfUse", "feedbackClarity", "confidence", "continuance"]
RATING_NUMERIC_COLUMNS = ["pitchScore", "rhythmScore"]
ANALYSIS_NUMERIC_COLUMNS = ["overallPitchScore", "overallRhythmScore", "confidence"]
VALIDATION_NUMERIC_COLUMNS = [
    "overallAgreement",
    "noteMatchedCount",
    "notePrecision",
    "noteRecall",
    "noteF1",
    "measureMatchedCount",
    "measurePrecision",
    "measureRecall",
    "measureF1",
]

PREPOST_TIME_ORDER = ["pretest", "posttest"]
PRACTICE_PATH_ORDER = ["pitch-first", "rhythm-first", "review-first"]
METRIC_SPECS = [
    ("pitch", "pretestPitch", "posttestPitch"),
    ("rhythm", "pretestRhythm", "posttestRhythm"),
]
ADJUDICATION_OVERALL_GAP_THRESHOLD = 2.0
ADJUDICATION_NOTE_F1_THRESHOLD = 0.67
ADJUDICATION_MEASURE_F1_THRESHOLD = 0.67
GROUP_LABELS = {"experimental": "实验组", "control": "对照组"}
METRIC_LABELS = {
    "pitch": "音准",
    "rhythm": "节奏",
    "pitchGain": "音准增益",
    "rhythmGain": "节奏增益",
    "usefulness": "感知有用性",
    "continuance": "持续使用意愿",
    "analysisCount": "分析使用次数",
}


def safe_read_csv(path: Path) -> pd.DataFrame:
    if not path.exists():
        raise FileNotFoundError(f"missing file: {path}")
    return pd.read_csv(path)


def safe_read_optional_csv(path: Path | None, columns: list[str] | None = None) -> pd.DataFrame:
    if path is None or not path.exists():
        return pd.DataFrame(columns=columns or [])
    return pd.read_csv(path)


def ensure_columns(frame: pd.DataFrame, columns: list[str]) -> pd.DataFrame:
    for column in columns:
        if column not in frame.columns:
            frame[column] = pd.NA
    return frame


def ensure_numeric(frame: pd.DataFrame, columns: list[str]) -> pd.DataFrame:
    frame = ensure_columns(frame, columns)
    for column in columns:
        frame[column] = pd.to_numeric(frame[column], errors="coerce")
    return frame


def save_table(frame: pd.DataFrame, output_dir: Path, name: str) -> None:
    frame.to_csv(output_dir / f"{name}.csv", index=False, encoding="utf-8-sig")


def save_figure(fig: plt.Figure, output_dir: Path, name: str) -> None:
    fig.tight_layout()
    fig.savefig(output_dir / f"{name}.png", dpi=220, bbox_inches="tight")
    plt.close(fig)


def round_value(value: float | int | None, digits: int = 3) -> float | None:
    if value is None or pd.isna(value):
        return None
    return round(float(value), digits)


def format_number(value: float | int | None, digits: int = 2, default: str = "未报告") -> str:
    if value is None or pd.isna(value):
        return default
    return f"{float(value):.{digits}f}"


def format_p_value(value: float | int | None) -> str:
    if value is None or pd.isna(value):
        return "未报告"
    numeric = float(value)
    if numeric < 0.001:
        return "< .001"
    return f"= {numeric:.3f}"


def p_value_interpretation(value: float | int | None) -> str:
    if value is None or pd.isna(value):
        return "当前样本不足，尚不能判断统计显著性"
    numeric = float(value)
    if numeric < 0.001:
        return "差异达到高度显著"
    if numeric < 0.01:
        return "差异达到显著"
    if numeric < 0.05:
        return "差异达到统计显著"
    return "差异未达到统计显著"


def find_row(frame: pd.DataFrame, **criteria) -> dict[str, object] | None:
    if frame.empty:
        return None
    subset = frame.copy()
    for key, expected in criteria.items():
        if key not in subset.columns:
            return None
        subset = subset.loc[subset[key] == expected]
    if subset.empty:
        return None
    return subset.iloc[0].to_dict()


def group_label(group_id: object) -> str:
    if pd.isna(group_id):
        return "未分组"
    return GROUP_LABELS.get(str(group_id), str(group_id))


def metric_label(metric: object) -> str:
    if pd.isna(metric):
        return "未命名指标"
    return METRIC_LABELS.get(str(metric), str(metric))


def safe_mean(series: pd.Series) -> float:
    cleaned = series.dropna()
    if cleaned.empty:
        return float("nan")
    return float(cleaned.mean())


def coerce_bool_series(series: pd.Series) -> pd.Series:
    def _coerce(value: object) -> bool | pd.NA:
        if pd.isna(value):
            return pd.NA
        if isinstance(value, bool):
            return value
        text = str(value).strip().lower()
        if text in {"true", "1", "yes", "y"}:
            return True
        if text in {"false", "0", "no", "n"}:
            return False
        return pd.NA

    return series.map(_coerce)


def safe_pearsonr(x: pd.Series, y: pd.Series) -> tuple[float, float, int]:
    subset = pd.DataFrame({"x": x, "y": y}).dropna()
    if len(subset) < 3 or subset["x"].nunique() < 2 or subset["y"].nunique() < 2:
        return float("nan"), float("nan"), len(subset)
    corr, p_value = stats.pearsonr(subset["x"], subset["y"])
    return float(corr), float(p_value), len(subset)


def placeholder_figure(title: str, message: str = "No data available") -> plt.Figure:
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.axis("off")
    ax.text(0.5, 0.62, title, ha="center", va="center", fontsize=15, fontweight="bold")
    ax.text(0.5, 0.45, message, ha="center", va="center", fontsize=11, color="#555555")
    return fig


def add_placeholder_axis(ax: plt.Axes, title: str, message: str = "No data available") -> None:
    ax.axis("off")
    ax.text(0.5, 0.6, title, ha="center", va="center", fontsize=13, fontweight="bold")
    ax.text(0.5, 0.42, message, ha="center", va="center", fontsize=10, color="#555555")


def group_summary(participants: pd.DataFrame) -> pd.DataFrame:
    participants = ensure_columns(
        participants,
        [
            "participantId",
            "groupId",
            "analysisCount",
            "pretestPitch",
            "posttestPitch",
            "pitchGain",
            "pretestRhythm",
            "posttestRhythm",
            "rhythmGain",
            "usefulness",
            "continuance",
            "questionnaireCount",
        ],
    )
    grouped = participants.groupby("groupId", dropna=False).agg(
        participantCount=("participantId", "count"),
        analysisCount=("analysisCount", "mean"),
        pretestPitchMean=("pretestPitch", "mean"),
        posttestPitchMean=("posttestPitch", "mean"),
        pitchGainMean=("pitchGain", "mean"),
        pretestRhythmMean=("pretestRhythm", "mean"),
        posttestRhythmMean=("posttestRhythm", "mean"),
        rhythmGainMean=("rhythmGain", "mean"),
        usefulnessMean=("usefulness", "mean"),
        continuanceMean=("continuance", "mean"),
        questionnaireCountMean=("questionnaireCount", "mean"),
    )
    return grouped.reset_index()


def ttest_summary(participants: pd.DataFrame) -> pd.DataFrame:
    participants = ensure_columns(participants, ["groupId", "pitchGain", "rhythmGain", "usefulness", "continuance", "analysisCount"])
    rows = []
    experimental = participants.loc[participants["groupId"] == "experimental"]
    control = participants.loc[participants["groupId"] == "control"]
    metrics = ["pitchGain", "rhythmGain", "usefulness", "continuance", "analysisCount"]
    for metric in metrics:
        exp_values = experimental[metric].dropna()
        ctl_values = control[metric].dropna()
        if len(exp_values) >= 2 and len(ctl_values) >= 2:
            statistic, p_value = stats.ttest_ind(exp_values, ctl_values, equal_var=False)
        else:
            statistic, p_value = float("nan"), float("nan")
        rows.append(
            {
                "metric": metric,
                "experimentalMean": safe_mean(exp_values),
                "controlMean": safe_mean(ctl_values),
                "tStatistic": statistic,
                "pValue": p_value,
            }
        )
    return pd.DataFrame(rows)


def questionnaire_summary(questionnaires: pd.DataFrame) -> pd.DataFrame:
    questionnaires = ensure_columns(questionnaires, ["groupId", "sessionStage", *QUESTIONNAIRE_NUMERIC_COLUMNS])
    grouped = questionnaires.groupby(["groupId", "sessionStage"], dropna=False)[QUESTIONNAIRE_NUMERIC_COLUMNS].mean().reset_index()
    return grouped.sort_values(["groupId", "sessionStage"])


def build_expert_system_table(participants: pd.DataFrame) -> pd.DataFrame:
    participants = ensure_columns(
        participants,
        ["pretestPitch", "expertPretestPitch", "posttestPitch", "expertPosttestPitch", "pretestRhythm", "expertPretestRhythm", "posttestRhythm", "expertPosttestRhythm"],
    )
    rows = []
    pairs = [
        ("pretestPitch", "expertPretestPitch"),
        ("posttestPitch", "expertPosttestPitch"),
        ("pretestRhythm", "expertPretestRhythm"),
        ("posttestRhythm", "expertPosttestRhythm"),
    ]
    for system_col, expert_col in pairs:
        corr, p_value, sample_size = safe_pearsonr(participants[system_col], participants[expert_col])
        rows.append(
            {
                "systemMetric": system_col,
                "expertMetric": expert_col,
                "sampleSize": sample_size,
                "pearsonR": corr,
                "pValue": p_value,
            }
        )
    return pd.DataFrame(rows)


def build_analysis_usage_table(analyses: pd.DataFrame) -> pd.DataFrame:
    analyses = ensure_columns(analyses, ["participantId", "sessionStage", "analysisId", "overallPitchScore", "overallRhythmScore", "confidence"])
    grouped = analyses.groupby(["participantId", "sessionStage"], dropna=False).agg(
        analysisRuns=("analysisId", "count"),
        meanPitchScore=("overallPitchScore", "mean"),
        meanRhythmScore=("overallRhythmScore", "mean"),
        meanConfidence=("confidence", "mean"),
    )
    return grouped.reset_index().sort_values(["participantId", "sessionStage"])


def build_usage_correlation_table(participants: pd.DataFrame) -> pd.DataFrame:
    participants = ensure_columns(participants, ["analysisCount", "pitchGain", "rhythmGain"])
    rows = []
    for metric in ["pitchGain", "rhythmGain"]:
        corr, p_value, sample_size = safe_pearsonr(participants["analysisCount"], participants[metric])
        rows.append(
            {
                "predictor": "analysisCount",
                "outcome": metric,
                "sampleSize": sample_size,
                "pearsonR": corr,
                "pValue": p_value,
            }
        )
    return pd.DataFrame(rows)


def prepare_validation_reviews(validation_reviews: pd.DataFrame) -> pd.DataFrame:
    validation_reviews = ensure_columns(
        validation_reviews,
        [
            "reviewId",
            "participantId",
            "groupId",
            "sessionStage",
            "teacherPrimaryPath",
            "systemRecommendedPath",
            "pathAgreement",
            "submittedAt",
            *VALIDATION_NUMERIC_COLUMNS,
        ],
    )
    validation_reviews = ensure_numeric(validation_reviews, VALIDATION_NUMERIC_COLUMNS)
    validation_reviews["pathAgreement"] = coerce_bool_series(validation_reviews["pathAgreement"])
    validation_reviews["teacherPrimaryPath"] = validation_reviews["teacherPrimaryPath"].fillna("review-first")
    validation_reviews["systemRecommendedPath"] = validation_reviews["systemRecommendedPath"].fillna("review-first")
    return validation_reviews


def build_validation_summary(validation_reviews: pd.DataFrame) -> pd.DataFrame:
    validation_reviews = prepare_validation_reviews(validation_reviews)
    if validation_reviews.empty:
        return pd.DataFrame(
            [
                {
                    "reviewCount": 0,
                    "participantCount": 0,
                    "averageAgreement": float("nan"),
                    "averageAgreementNormalized": float("nan"),
                    "averageNotePrecision": float("nan"),
                    "averageNoteRecall": float("nan"),
                    "averageNoteF1": float("nan"),
                    "averageMeasurePrecision": float("nan"),
                    "averageMeasureRecall": float("nan"),
                    "averageMeasureF1": float("nan"),
                    "pathAgreementRate": float("nan"),
                }
            ]
        )

    agreement_numeric = validation_reviews["pathAgreement"].map(
        lambda value: float(value) if pd.notna(value) else float("nan"),
    )
    return pd.DataFrame(
        [
            {
                "reviewCount": len(validation_reviews),
                "participantCount": validation_reviews["participantId"].dropna().nunique(),
                "averageAgreement": validation_reviews["overallAgreement"].mean(),
                "averageAgreementNormalized": validation_reviews["overallAgreement"].mean() / 5 if validation_reviews["overallAgreement"].notna().any() else float("nan"),
                "averageNotePrecision": validation_reviews["notePrecision"].mean(),
                "averageNoteRecall": validation_reviews["noteRecall"].mean(),
                "averageNoteF1": validation_reviews["noteF1"].mean(),
                "averageMeasurePrecision": validation_reviews["measurePrecision"].mean(),
                "averageMeasureRecall": validation_reviews["measureRecall"].mean(),
                "averageMeasureF1": validation_reviews["measureF1"].mean(),
                "pathAgreementRate": agreement_numeric.mean(),
            }
        ]
    )


def build_validation_group_summary(validation_reviews: pd.DataFrame) -> pd.DataFrame:
    validation_reviews = prepare_validation_reviews(validation_reviews)
    if validation_reviews.empty:
        return pd.DataFrame(
            columns=[
                "groupId",
                "reviewCount",
                "participantCount",
                "averageAgreement",
                "averageAgreementNormalized",
                "averageNoteF1",
                "averageMeasureF1",
                "pathAgreementRate",
            ]
        )

    review_counts = validation_reviews.groupby("groupId", dropna=False)["reviewId"].count().rename("reviewCount")
    participant_counts = validation_reviews.groupby("groupId", dropna=False)["participantId"].nunique().rename("participantCount")
    grouped = validation_reviews.groupby("groupId", dropna=False).agg(
        averageAgreement=("overallAgreement", "mean"),
        averageNoteF1=("noteF1", "mean"),
        averageMeasureF1=("measureF1", "mean"),
    )
    agreement_rate = (
        validation_reviews.assign(pathAgreementNumeric=validation_reviews["pathAgreement"].map(lambda value: float(value) if pd.notna(value) else float("nan")))
        .groupby("groupId", dropna=False)["pathAgreementNumeric"]
        .mean()
        .rename("pathAgreementRate")
    )
    summary = pd.concat([review_counts, participant_counts, grouped, agreement_rate], axis=1).reset_index()
    summary["averageAgreementNormalized"] = summary["averageAgreement"] / 5
    return summary.sort_values("groupId")


def build_validation_path_confusion(validation_reviews: pd.DataFrame) -> pd.DataFrame:
    validation_reviews = prepare_validation_reviews(validation_reviews)
    if validation_reviews.empty:
        return pd.DataFrame(columns=["teacherPrimaryPath", "systemRecommendedPath", "reviewCount"])

    confusion = (
        validation_reviews.groupby(["teacherPrimaryPath", "systemRecommendedPath"], dropna=False)
        .size()
        .rename("reviewCount")
        .reset_index()
    )
    return confusion.sort_values(["teacherPrimaryPath", "systemRecommendedPath"])


def parse_pipe_set(value: object, numeric: bool = False) -> set[object]:
    if pd.isna(value):
        return set()
    items = [item.strip() for item in str(value).split("|") if item.strip()]
    if numeric:
        parsed: set[object] = set()
        for item in items:
            try:
                parsed.add(int(float(item)))
            except ValueError:
                continue
        return parsed
    return set(items)


def compute_set_f1(left: set[object], right: set[object]) -> tuple[float | None, float | None, float | None]:
    if not left and not right:
        return None, None, None
    matched = len(left & right)
    precision = matched / len(left) if left else None
    recall = matched / len(right) if right else None
    if precision is None or recall is None or (precision + recall) == 0:
        f1 = None
    else:
        f1 = (2 * precision * recall) / (precision + recall)
    return precision, recall, f1


def cohen_kappa_score(left: pd.Series, right: pd.Series, categories: list[str]) -> float:
    subset = pd.DataFrame({"left": left, "right": right}).dropna()
    if subset.empty:
        return float("nan")
    observed = float((subset["left"] == subset["right"]).mean())
    expected = 0.0
    for category in categories:
        expected += float((subset["left"] == category).mean()) * float((subset["right"] == category).mean())
    if expected >= 1.0:
        return float("nan")
    return (observed - expected) / (1.0 - expected)


def icc_2_1(frame: pd.DataFrame) -> float:
    matrix = frame.dropna().to_numpy(dtype=float)
    if matrix.ndim != 2 or matrix.shape[0] < 2 or matrix.shape[1] < 2:
        return float("nan")
    n, k = matrix.shape
    grand_mean = matrix.mean()
    subject_means = matrix.mean(axis=1)
    rater_means = matrix.mean(axis=0)
    ss_subject = k * ((subject_means - grand_mean) ** 2).sum()
    ss_rater = n * ((rater_means - grand_mean) ** 2).sum()
    ss_error = ((matrix - subject_means[:, None] - rater_means[None, :] + grand_mean) ** 2).sum()
    ms_subject = ss_subject / (n - 1)
    ms_rater = ss_rater / (k - 1)
    ms_error = ss_error / ((n - 1) * (k - 1))
    denominator = ms_subject + (k - 1) * ms_error + (k * (ms_rater - ms_error) / n)
    if denominator == 0:
        return float("nan")
    return float((ms_subject - ms_error) / denominator)


def adjudication_reasons(row: dict[str, object]) -> list[str]:
    reasons: list[str] = []
    if not bool(row.get("pathMatch", False)):
        reasons.append("practice-path mismatch")
    overall_gap = row.get("overallAgreementGap")
    if overall_gap is not None and not pd.isna(overall_gap) and float(overall_gap) >= ADJUDICATION_OVERALL_GAP_THRESHOLD:
        reasons.append("overall-agreement gap >= 2")
    note_overlap = row.get("noteOverlapF1")
    if note_overlap is not None and not pd.isna(note_overlap) and float(note_overlap) < ADJUDICATION_NOTE_F1_THRESHOLD:
        reasons.append("note-overlap F1 < 0.67")
    measure_overlap = row.get("measureOverlapF1")
    if measure_overlap is not None and not pd.isna(measure_overlap) and float(measure_overlap) < ADJUDICATION_MEASURE_F1_THRESHOLD:
        reasons.append("measure-overlap F1 < 0.67")
    return reasons


def build_inter_rater_pairs(validation_reviews: pd.DataFrame) -> pd.DataFrame:
    validation_reviews = prepare_validation_reviews(validation_reviews)
    validation_reviews = ensure_columns(
        validation_reviews,
        [
            "analysisId",
            "participantId",
            "groupId",
            "sessionStage",
            "pieceId",
            "sectionId",
            "raterId",
            "overallAgreement",
            "teacherPrimaryPath",
            "teacherIssueNoteIds",
            "teacherIssueMeasureIndexes",
            "submittedAt",
        ],
    )
    if validation_reviews.empty:
        return pd.DataFrame(
            columns=[
                "analysisId",
                "participantId",
                "groupId",
                "sessionStage",
                "pieceId",
                "sectionId",
                "scoreUnit",
                "raterAId",
                "raterBId",
                "overallAgreementA",
                "overallAgreementB",
                "overallAgreementGap",
                "teacherPrimaryPathA",
                "teacherPrimaryPathB",
                "pathMatch",
                "noteOverlapPrecision",
                "noteOverlapRecall",
                "noteOverlapF1",
                "measureOverlapPrecision",
                "measureOverlapRecall",
                "measureOverlapF1",
                "requiresAdjudication",
                "adjudicationReason",
            ]
        )

    reviews = validation_reviews.sort_values(["analysisId", "submittedAt", "raterId"]).drop_duplicates(["analysisId", "raterId"], keep="last")
    rows: list[dict[str, object]] = []
    for analysis_id, analysis_frame in reviews.groupby("analysisId", dropna=False):
        analysis_rows = analysis_frame.sort_values(["submittedAt", "raterId"]).to_dict(orient="records")
        if len(analysis_rows) < 2:
            continue
        first = analysis_rows[0]
        second = analysis_rows[1]
        note_precision, note_recall, note_f1 = compute_set_f1(
            parse_pipe_set(first.get("teacherIssueNoteIds")),
            parse_pipe_set(second.get("teacherIssueNoteIds")),
        )
        measure_precision, measure_recall, measure_f1 = compute_set_f1(
            parse_pipe_set(first.get("teacherIssueMeasureIndexes"), numeric=True),
            parse_pipe_set(second.get("teacherIssueMeasureIndexes"), numeric=True),
        )
        rows.append(
            {
                "analysisId": analysis_id,
                "participantId": first.get("participantId"),
                "groupId": first.get("groupId"),
                "sessionStage": first.get("sessionStage"),
                "pieceId": first.get("pieceId"),
                "sectionId": first.get("sectionId"),
                "scoreUnit": f"{first.get('pieceId')}/{first.get('sectionId')}",
                "raterAId": first.get("raterId"),
                "raterBId": second.get("raterId"),
                "overallAgreementA": first.get("overallAgreement"),
                "overallAgreementB": second.get("overallAgreement"),
                "overallAgreementGap": abs(float(first.get("overallAgreement") or 0) - float(second.get("overallAgreement") or 0)),
                "teacherPrimaryPathA": first.get("teacherPrimaryPath"),
                "teacherPrimaryPathB": second.get("teacherPrimaryPath"),
                "pathMatch": first.get("teacherPrimaryPath") == second.get("teacherPrimaryPath"),
                "noteOverlapPrecision": note_precision,
                "noteOverlapRecall": note_recall,
                "noteOverlapF1": note_f1,
                "measureOverlapPrecision": measure_precision,
                "measureOverlapRecall": measure_recall,
                "measureOverlapF1": measure_f1,
            }
        )
    pairs = pd.DataFrame(rows)
    if pairs.empty:
        return pairs
    pairs["adjudicationReason"] = pairs.apply(lambda row: " | ".join(adjudication_reasons(row.to_dict())), axis=1)
    pairs["requiresAdjudication"] = pairs["adjudicationReason"].astype(str).str.len().gt(0)
    return pairs


def build_inter_rater_summary(inter_rater_pairs: pd.DataFrame) -> pd.DataFrame:
    inter_rater_pairs = ensure_columns(
        inter_rater_pairs,
        [
            "analysisId",
            "participantId",
            "overallAgreementA",
            "overallAgreementB",
            "overallAgreementGap",
            "teacherPrimaryPathA",
            "teacherPrimaryPathB",
            "pathMatch",
            "noteOverlapF1",
            "measureOverlapF1",
            "requiresAdjudication",
        ],
    )
    if inter_rater_pairs.empty:
        return pd.DataFrame(
            [
                {
                    "pairCount": 0,
                    "analysisCount": 0,
                    "participantCount": 0,
                    "pathCohenKappa": float("nan"),
                    "overallAgreementICC": float("nan"),
                    "pathMatchRate": float("nan"),
                    "meanAgreementGap": float("nan"),
                    "meanNoteOverlapF1": float("nan"),
                    "meanMeasureOverlapF1": float("nan"),
                    "adjudicationCount": 0,
                    "adjudicationRate": float("nan"),
                }
            ]
        )

    icc_value = icc_2_1(inter_rater_pairs[["overallAgreementA", "overallAgreementB"]])
    kappa_value = cohen_kappa_score(
        inter_rater_pairs["teacherPrimaryPathA"],
        inter_rater_pairs["teacherPrimaryPathB"],
        PRACTICE_PATH_ORDER,
    )
    return pd.DataFrame(
        [
            {
                "pairCount": len(inter_rater_pairs),
                "analysisCount": inter_rater_pairs["analysisId"].nunique(),
                "participantCount": inter_rater_pairs["participantId"].nunique(),
                "pathCohenKappa": kappa_value,
                "overallAgreementICC": icc_value,
                "pathMatchRate": inter_rater_pairs["pathMatch"].map(lambda value: float(bool(value)) if pd.notna(value) else float("nan")).mean(),
                "meanAgreementGap": inter_rater_pairs["overallAgreementGap"].mean(),
                "meanNoteOverlapF1": inter_rater_pairs["noteOverlapF1"].mean(),
                "meanMeasureOverlapF1": inter_rater_pairs["measureOverlapF1"].mean(),
                "adjudicationCount": int(inter_rater_pairs["requiresAdjudication"].fillna(False).astype(bool).sum()),
                "adjudicationRate": inter_rater_pairs["requiresAdjudication"].map(lambda value: float(bool(value)) if pd.notna(value) else float("nan")).mean(),
            }
        ]
    )


def build_inter_rater_breakdown(inter_rater_pairs: pd.DataFrame, group_columns: list[str]) -> pd.DataFrame:
    required_columns = group_columns + [
        "analysisId",
        "participantId",
        "overallAgreementA",
        "overallAgreementB",
        "overallAgreementGap",
        "teacherPrimaryPathA",
        "teacherPrimaryPathB",
        "pathMatch",
        "noteOverlapF1",
        "measureOverlapF1",
        "requiresAdjudication",
    ]
    inter_rater_pairs = ensure_columns(inter_rater_pairs, required_columns)
    if inter_rater_pairs.empty:
        return pd.DataFrame(
            columns=group_columns
            + [
                "pairCount",
                "analysisCount",
                "participantCount",
                "pathCohenKappa",
                "overallAgreementICC",
                "pathMatchRate",
                "meanAgreementGap",
                "meanNoteOverlapF1",
                "meanMeasureOverlapF1",
                "adjudicationCount",
                "adjudicationRate",
            ]
        )

    rows: list[dict[str, object]] = []
    group_source: str | list[str] = group_columns[0] if len(group_columns) == 1 else group_columns
    for group_key, frame in inter_rater_pairs.groupby(group_source, dropna=False):
        summary = build_inter_rater_summary(frame).iloc[0].to_dict()
        row: dict[str, object] = {}
        if len(group_columns) == 1:
            row[group_columns[0]] = group_key
        else:
            for column, value in zip(group_columns, group_key):
                row[column] = value
        row.update(summary)
        rows.append(row)

    return pd.DataFrame(rows).sort_values(group_columns)


def build_inter_rater_adjudication_queue(inter_rater_pairs: pd.DataFrame) -> pd.DataFrame:
    inter_rater_pairs = ensure_columns(
        inter_rater_pairs,
        [
            "analysisId",
            "participantId",
            "groupId",
            "sessionStage",
            "pieceId",
            "sectionId",
            "scoreUnit",
            "raterAId",
            "raterBId",
            "overallAgreementGap",
            "pathMatch",
            "noteOverlapF1",
            "measureOverlapF1",
            "requiresAdjudication",
            "adjudicationReason",
        ],
    )
    if inter_rater_pairs.empty:
        return inter_rater_pairs.iloc[0:0].copy()
    queue = inter_rater_pairs.loc[inter_rater_pairs["requiresAdjudication"].fillna(False).astype(bool)].copy()
    if queue.empty:
        return queue
    return queue.sort_values(["sessionStage", "pieceId", "sectionId", "analysisId"])


def build_prepost_long_table(participants: pd.DataFrame) -> pd.DataFrame:
    participants = ensure_columns(participants, ["participantId", "groupId", "pretestPitch", "posttestPitch", "pretestRhythm", "posttestRhythm"])
    rows: list[dict[str, object]] = []
    for metric, pre_col, post_col in METRIC_SPECS:
        subset = participants[["participantId", "groupId", pre_col, post_col]].copy()
        for _, row in subset.iterrows():
            if pd.notna(row[pre_col]):
                rows.append(
                    {
                        "participantId": row["participantId"],
                        "groupId": row["groupId"],
                        "metric": metric,
                        "time": "pretest",
                        "score": row[pre_col],
                    }
                )
            if pd.notna(row[post_col]):
                rows.append(
                    {
                        "participantId": row["participantId"],
                        "groupId": row["groupId"],
                        "metric": metric,
                        "time": "posttest",
                        "score": row[post_col],
                    }
                )
    long_table = pd.DataFrame(rows, columns=["participantId", "groupId", "metric", "time", "score"])
    if not long_table.empty:
        long_table["time"] = pd.Categorical(long_table["time"], categories=PREPOST_TIME_ORDER, ordered=True)
        long_table = long_table.sort_values(["metric", "groupId", "participantId", "time"])
    return long_table


def build_prepost_summary(long_table: pd.DataFrame) -> pd.DataFrame:
    long_table = ensure_columns(long_table, ["metric", "groupId", "time", "score"])
    if long_table.empty:
        return pd.DataFrame(columns=["metric", "groupId", "time", "sampleSize", "meanScore", "sd", "sem"])
    grouped = long_table.groupby(["metric", "groupId", "time"], dropna=False).agg(
        sampleSize=("score", "count"),
        meanScore=("score", "mean"),
        sd=("score", "std"),
        sem=("score", "sem"),
    )
    summary = grouped.reset_index()
    summary["time"] = pd.Categorical(summary["time"], categories=PREPOST_TIME_ORDER, ordered=True)
    return summary.sort_values(["metric", "groupId", "time"])


def build_ancova_table(participants: pd.DataFrame) -> pd.DataFrame:
    participants = ensure_columns(participants, ["participantId", "groupId", "pretestPitch", "posttestPitch", "pretestRhythm", "posttestRhythm"])
    rows = []
    for metric, pre_col, post_col in METRIC_SPECS:
        subset = participants[["participantId", "groupId", pre_col, post_col]].dropna().copy()
        subset["groupId"] = subset["groupId"].astype("string")
        if len(subset) < 4 or subset["groupId"].nunique() < 2:
            for term in ["groupId", "pretest"]:
                rows.append(
                    {
                        "metric": metric,
                        "term": term,
                        "sampleSize": len(subset),
                        "degreesOfFreedom": float("nan"),
                        "sumSquares": float("nan"),
                        "fStatistic": float("nan"),
                        "pValue": float("nan"),
                        "partialEtaSquared": float("nan"),
                        "adjustedRsquared": float("nan"),
                    }
                )
            continue

        try:
            model = ols(f"{post_col} ~ C(groupId) + {pre_col}", data=subset).fit()
            anova_table = anova_lm(model, typ=2)
            residual_ss = float(anova_table.loc["Residual", "sum_sq"]) if "Residual" in anova_table.index else float("nan")
            adjusted_rsq = float(model.rsquared_adj)
            term_lookup = {
                "C(groupId)": "groupId",
                pre_col: "pretest",
            }

            for source_name, term in term_lookup.items():
                if source_name not in anova_table.index:
                    rows.append(
                        {
                            "metric": metric,
                            "term": term,
                            "sampleSize": len(subset),
                            "degreesOfFreedom": float("nan"),
                            "sumSquares": float("nan"),
                            "fStatistic": float("nan"),
                            "pValue": float("nan"),
                            "partialEtaSquared": float("nan"),
                            "adjustedRsquared": adjusted_rsq,
                        }
                    )
                    continue

                sum_squares = float(anova_table.loc[source_name, "sum_sq"])
                partial_eta = float("nan")
                if pd.notna(residual_ss) and (sum_squares + residual_ss) > 0:
                    partial_eta = sum_squares / (sum_squares + residual_ss)
                rows.append(
                    {
                        "metric": metric,
                        "term": term,
                        "sampleSize": len(subset),
                        "degreesOfFreedom": float(anova_table.loc[source_name, "df"]),
                        "sumSquares": sum_squares,
                        "fStatistic": float(anova_table.loc[source_name, "F"]),
                        "pValue": float(anova_table.loc[source_name, "PR(>F)"]),
                        "partialEtaSquared": partial_eta,
                        "adjustedRsquared": adjusted_rsq,
                    }
                )
        except Exception:
            for term in ["groupId", "pretest"]:
                rows.append(
                    {
                        "metric": metric,
                        "term": term,
                        "sampleSize": len(subset),
                        "degreesOfFreedom": float("nan"),
                        "sumSquares": float("nan"),
                        "fStatistic": float("nan"),
                        "pValue": float("nan"),
                        "partialEtaSquared": float("nan"),
                        "adjustedRsquared": float("nan"),
                    }
                )
    return pd.DataFrame(rows)


def build_prepost_sentence(metric: str, prepost_summary: pd.DataFrame) -> str:
    exp_pre = find_row(prepost_summary, metric=metric, groupId="experimental", time="pretest")
    exp_post = find_row(prepost_summary, metric=metric, groupId="experimental", time="posttest")
    ctl_pre = find_row(prepost_summary, metric=metric, groupId="control", time="pretest")
    ctl_post = find_row(prepost_summary, metric=metric, groupId="control", time="posttest")
    label = metric_label(metric)
    if not all([exp_pre, exp_post, ctl_pre, ctl_post]):
        return f"{label}前后测数据尚未完整导入，当前结果段落保留为写作模板。"

    return (
        f"在{label}指标上，实验组前测均值为 {format_number(exp_pre.get('meanScore'))} "
        f"(n = {int(exp_pre.get('sampleSize', 0))})，后测均值提升至 {format_number(exp_post.get('meanScore'))} "
        f"(n = {int(exp_post.get('sampleSize', 0))})；对照组前测均值为 {format_number(ctl_pre.get('meanScore'))} "
        f"(n = {int(ctl_pre.get('sampleSize', 0))})，后测均值为 {format_number(ctl_post.get('meanScore'))} "
        f"(n = {int(ctl_post.get('sampleSize', 0))})。"
    )


def build_ttest_sentence(metric: str, ttest_table: pd.DataFrame) -> str:
    row = find_row(ttest_table, metric=metric)
    label = metric_label(metric)
    if row is None:
        return f"{label}的组间增益比较尚未生成。"
    return (
        f"{label}的组间比较显示，实验组均值为 {format_number(row.get('experimentalMean'))}，"
        f"对照组均值为 {format_number(row.get('controlMean'))}，"
        f"Welch t = {format_number(row.get('tStatistic'), 3)}，p {format_p_value(row.get('pValue'))}，"
        f"{p_value_interpretation(row.get('pValue'))}。"
    )


def build_ancova_sentence(metric: str, ancova_table: pd.DataFrame) -> str:
    row = find_row(ancova_table, metric=metric, term="groupId")
    label = metric_label(metric)
    if row is None:
        return f"{label}的 ANCOVA 结果尚未生成。"
    return (
        f"以对应前测成绩为协变量后，组别对后测{label}的主效应为 "
        f"F = {format_number(row.get('fStatistic'), 3)}，p {format_p_value(row.get('pValue'))}，"
        f"partial η² = {format_number(row.get('partialEtaSquared'), 3)}，"
        f"调整后 R² = {format_number(row.get('adjustedRsquared'), 3)}。"
    )


def build_expert_sentence(metric: str, expert_table: pd.DataFrame) -> str:
    metric_map = {
        "pitch": ("posttestPitch", "expertPosttestPitch"),
        "rhythm": ("posttestRhythm", "expertPosttestRhythm"),
    }
    if metric not in metric_map:
        return ""
    system_metric, expert_metric = metric_map[metric]
    row = find_row(expert_table, systemMetric=system_metric, expertMetric=expert_metric)
    label = metric_label(metric)
    if row is None:
        return f"系统与专家在后测{label}上的一致性结果尚未生成。"
    return (
        f"系统后测{label}得分与专家后测{label}评分的相关系数为 "
        f"r = {format_number(row.get('pearsonR'), 3)}，p {format_p_value(row.get('pValue'))}，"
        f"样本量为 n = {int(row.get('sampleSize', 0))}。"
    )


def build_usage_sentence(metric: str, usage_corr_table: pd.DataFrame) -> str:
    outcome = f"{metric}Gain"
    row = find_row(usage_corr_table, predictor="analysisCount", outcome=outcome)
    label = metric_label(outcome)
    if row is None:
        return f"分析使用次数与{label}之间的相关结果尚未生成。"
    return (
        f"分析使用次数与{label}之间的相关系数为 "
        f"r = {format_number(row.get('pearsonR'), 3)}，p {format_p_value(row.get('pValue'))}，"
        f"样本量为 n = {int(row.get('sampleSize', 0))}。"
    )


def build_validation_sentence(validation_summary: pd.DataFrame, validation_group_summary: pd.DataFrame) -> str:
    if validation_summary.empty:
        return "鏁欏笀鏍囨敞楠岃瘉鏁版嵁灏氭湭瀵煎叆锛屽綋鍓嶆棤娉曟姤鍛婄郴缁熻緭鍑轰笌鏁欏笀鍒ゆ柇鐨勪竴鑷存€с€?"

    row = validation_summary.iloc[0].to_dict()
    review_count = int(row.get("reviewCount", 0) or 0)
    participant_count = int(row.get("participantCount", 0) or 0)
    if review_count == 0:
        return "鏁欏笀鏍囨敞楠岃瘉鏁版嵁灏氭湭瀵煎叆锛屽綋鍓嶆棤娉曟姤鍛婄郴缁熻緭鍑轰笌鏁欏笀鍒ゆ柇鐨勪竴鑷存€с€?"

    sentence = (
        f"鍦ㄦ暀甯堟爣娉ㄩ獙璇佸瓙鏍锋湰涓紝鍏辩撼鍏?{review_count} 鏉￠獙璇佽褰曪紝瑕嗙洊 {participant_count} 鍚嶅彈璇曡€咃紱"
        f"绯荤粺涓庢暀甯堢殑鏁翠綋涓€鑷存€у潎鍊间负 {format_number(row.get('averageAgreement'))}/5锛?"
        f"闊崇绾?F1 涓?{format_number(row.get('averageNoteF1'), 3)}锛?"
        f"灏忚妭绾?F1 涓?{format_number(row.get('averageMeasureF1'), 3)}锛?"
        f"缁冧範璺緞涓€鑷寸巼涓?{format_number((row.get('pathAgreementRate') or 0) * 100, 1)}%銆?"
    )

    if not validation_group_summary.empty:
        exp = find_row(validation_group_summary, groupId="experimental")
        ctl = find_row(validation_group_summary, groupId="control")
        if exp and ctl:
            sentence += (
                f"鍏朵腑瀹為獙缁勮矾寰勪竴鑷寸巼涓?{format_number((exp.get('pathAgreementRate') or 0) * 100, 1)}%锛?"
                f"瀵圭収缁勪负 {format_number((ctl.get('pathAgreementRate') or 0) * 100, 1)}%銆?"
            )
    return sentence


def build_inter_rater_sentence(inter_rater_summary: pd.DataFrame) -> str:
    if inter_rater_summary.empty:
        return "褰撳墠灏氭棤鍙敤鐨勫弻璇勬暀甯堟暟鎹紝鍥犳鏃犳硶鎶ュ憡鏁欏笀闂翠竴鑷存€с€?"

    row = inter_rater_summary.iloc[0].to_dict()
    pair_count = int(row.get("pairCount", 0) or 0)
    if pair_count == 0:
        return "褰撳墠灏氭棤鍙敤鐨勫弻璇勬暀甯堟暟鎹紝鍥犳鏃犳硶鎶ュ憡鏁欏笀闂翠竴鑷存€с€?"

    return (
        f"鍦ㄥ弻璇勬暀甯堝瓙鏍锋湰涓紝鍏辩撼鍏?{pair_count} 瀵瑰弻璇勮褰曪紱"
        f"缁冧範璺緞鐨?Cohen's kappa 涓?{format_number(row.get('pathCohenKappa'), 3)}锛?"
        f"鏁翠綋涓€鑷存€ц瘎鍒嗙殑 ICC 涓?{format_number(row.get('overallAgreementICC'), 3)}锛?"
        f"鏁欏笀闂村闂闊崇殑骞冲潎閲嶅彔 F1 涓?{format_number(row.get('meanNoteOverlapF1'), 3)}锛?"
        f"瀵归棶棰樺皬鑺傜殑骞冲潎閲嶅彔 F1 涓?{format_number(row.get('meanMeasureOverlapF1'), 3)}銆?"
    )


def build_experience_sentence(group_table: pd.DataFrame, ttest_table: pd.DataFrame) -> str:
    exp_group = find_row(group_table, groupId="experimental") or {}
    ctl_group = find_row(group_table, groupId="control") or {}
    usefulness_test = find_row(ttest_table, metric="usefulness") or {}
    continuance_test = find_row(ttest_table, metric="continuance") or {}
    return (
        f"在学习体验方面，实验组感知有用性均值为 {format_number(exp_group.get('usefulnessMean'))}，"
        f"对照组为 {format_number(ctl_group.get('usefulnessMean'))}；"
        f"持续使用意愿方面，实验组均值为 {format_number(exp_group.get('continuanceMean'))}，"
        f"对照组为 {format_number(ctl_group.get('continuanceMean'))}。"
        f"其中感知有用性的组间比较 p {format_p_value(usefulness_test.get('pValue'))}，"
        f"持续使用意愿的组间比较 p {format_p_value(continuance_test.get('pValue'))}。"
    )


def build_main_findings(ttest_table: pd.DataFrame, ancova_table: pd.DataFrame) -> str:
    significant_findings: list[str] = []
    for metric in ["pitchGain", "rhythmGain"]:
        row = find_row(ttest_table, metric=metric)
        if row and pd.notna(row.get("pValue")) and float(row["pValue"]) < 0.05:
            significant_findings.append(f"{metric_label(metric)}达到组间显著差异")
    for metric in ["pitch", "rhythm"]:
        row = find_row(ancova_table, metric=metric, term="groupId")
        if row and pd.notna(row.get("pValue")) and float(row["pValue"]) < 0.05:
            significant_findings.append(f"{metric_label(metric)}的 ANCOVA 组别主效应显著")

    if not significant_findings:
        return "当前样本下尚未观察到明确的显著主效应，论文可先按模板撰写，待正式数据进入后自动更新。"
    return "；".join(significant_findings) + "。"


def build_paper_sections(
    participants: pd.DataFrame,
    group_table: pd.DataFrame,
    ttest_table: pd.DataFrame,
    expert_table: pd.DataFrame,
    usage_corr_table: pd.DataFrame,
    ancova_table: pd.DataFrame,
    prepost_summary: pd.DataFrame,
    validation_summary: pd.DataFrame,
    validation_group_summary: pd.DataFrame,
    inter_rater_summary: pd.DataFrame,
) -> tuple[str, list[tuple[str, list[str]]], str]:
    participants = ensure_columns(participants, ["participantId", "groupId", "analysisCount", "weeklySessionCount"])
    participant_count = len(participants)
    experimental_count = int((participants["groupId"] == "experimental").sum())
    control_count = int((participants["groupId"] == "control").sum())
    analysis_mean = format_number(participants["analysisCount"].mean())
    weekly_mean = format_number(participants["weeklySessionCount"].mean())
    title = "基于深度学习反馈的二胡 AI 教学干预研究论文草稿"

    abstract_lines = [
        "本研究旨在评估一套基于深度学习的二胡 AI 反馈工具，是否能够在高校器乐训练情境下提升学习者的音准、节奏与学习体验。",
        (
            f"研究采用前测-后测准实验设计，当前导入受试者共 {participant_count} 名，"
            f"其中实验组 {experimental_count} 名，对照组 {control_count} 名。"
            "实验组使用结合 torchcrepe、librosa 与规则诊断的非实时反馈原型进行练习，对照组接受常规练习安排。"
        ),
        (
            "结果分析包括前后测描述统计、Welch t 检验、以前测为协变量的 ANCOVA、"
            "系统与专家评分相关分析以及使用次数与学习增益的相关分析。"
        ),
        build_validation_sentence(validation_summary, validation_group_summary),
        build_main_findings(ttest_table, ancova_table),
        build_inter_rater_sentence(inter_rater_summary),
        "研究结论部分建议结合定量结果与访谈资料，讨论 AI 反馈在器乐练习中的教学价值、接受度与适用边界。",
    ]

    keyword_line = "二胡；人工智能；深度学习；音乐教育；教学干预；音准反馈"
    background_lines = [
        "在音乐教育研究中，AI 的价值不只体现在模型精度，更体现在其是否能够提供可被学习者理解和采纳的形成性反馈。",
        "针对二胡训练，本研究将系统分析限制在音准与节奏两个可量化维度，以降低系统复杂度并提升教育干预研究的可控性。",
        "与强调实时识别或商业化功能的工程路线不同，本研究关注的是：深度学习驱动的非实时反馈工具，是否能够支持大学生的器乐练习成效与学习体验。",
    ]
    method_lines = [
        (
            f"本研究采用前测-后测准实验设计，按班级或工作室进行分组，当前样本共 {participant_count} 名，"
            f"实验组 {experimental_count} 名，对照组 {control_count} 名。"
        ),
        (
            "实验组在 6 至 8 周干预期内使用本研究开发的 PWA 原型系统完成任务化练习。"
            "系统通过录音上传、乐谱对齐、逐音音高检测和节奏分析，对问题小节和问题音进行标记，并提供标准示范回放。"
        ),
        (
            "分析引擎采用预训练深度学习与规则诊断结合的路线："
            "torchcrepe/CREPE 负责逐帧 F0 估计，librosa 负责 onset 与节奏特征提取，"
            "MusicXML/MIDI 与 DTW 负责标准乐谱对齐，最终输出逐音与逐小节反馈。"
        ),
        (
            "测量指标包括系统音准/节奏得分、专家评分、问卷中的感知有用性、易用性、反馈清晰度、学习信心与持续使用意愿，"
            "并记录每位受试者的分析次数、周任务与访谈数据。"
        ),
        (
            f"当前数据中，受试者平均分析使用次数为 {analysis_mean} 次，平均周练习记录数为 {weekly_mean} 次。"
            "统计分析采用描述统计、Welch t 检验、ANCOVA 与皮尔逊相关分析。"
        ),
    ]
    results_lines = [
        (
            f"样本描述显示，当前数据库共纳入 {participant_count} 名受试者，"
            f"其中实验组 {experimental_count} 名，对照组 {control_count} 名。"
        ),
        build_prepost_sentence("pitch", prepost_summary),
        build_prepost_sentence("rhythm", prepost_summary),
        build_ttest_sentence("pitchGain", ttest_table),
        build_ttest_sentence("rhythmGain", ttest_table),
        build_ancova_sentence("pitch", ancova_table),
        build_ancova_sentence("rhythm", ancova_table),
        build_expert_sentence("pitch", expert_table),
        build_expert_sentence("rhythm", expert_table),
        build_validation_sentence(validation_summary, validation_group_summary),
        build_inter_rater_sentence(inter_rater_summary),
        build_usage_sentence("pitch", usage_corr_table),
        build_usage_sentence("rhythm", usage_corr_table),
        build_experience_sentence(group_table, ttest_table),
    ]
    discussion_lines = [
        "如果正式数据继续呈现实验组在音准或节奏增益上的优势，则可说明深度学习驱动的诊断反馈具有形成性教学价值。",
        "若系统与专家评分保持中高程度相关，则说明该工具可作为教师评分之外的辅助判断来源，而非替代教师。",
        "若使用次数与增益呈正相关，可进一步支持“持续练习中的反馈暴露强度”是影响学习结果的重要机制。",
        "若统计结果未达到显著，也可讨论样本量、干预周期、受试者基础水平差异、二胡表达性滑音对音高评估的影响等限制。"
        "这类解释仍然符合 SSCI 对教育技术干预研究的叙事方式。"
    ]
    conclusion_lines = [
        "本研究已经形成一条可复现的研究流程：AI 原型采集练习数据，统计脚本自动导出论文图表与文本草稿，研究者据此补充理论框架与质性解释。",
        "后续正式写作时，可在本自动草稿基础上补入文献综述、理论模型、伦理说明与访谈引文，以形成完整投稿稿件。",
    ]
    appendix_lines = [
        "建议在 Word 中保留以下附件材料：教师评分 rubric、访谈提纲、受试说明书、周任务模板、原始导出表与图表清单。",
        "如需直接写“结果”章节，可优先使用本次自动生成的 `results_section_zh.txt` 与 `paper_draft_zh.docx`。",
    ]

    sections = [
        ("摘要", abstract_lines),
        ("关键词", [keyword_line]),
        ("一、研究背景与问题提出", background_lines),
        ("二、研究设计与方法", method_lines),
        ("三、研究结果", results_lines),
        ("四、讨论", discussion_lines),
        ("五、结论", conclusion_lines),
        ("附录写作提示", appendix_lines),
    ]
    results_text = "\n\n".join(results_lines)
    return title, sections, results_text


def write_paper_draft(
    output_dir: Path,
    participants: pd.DataFrame,
    group_table: pd.DataFrame,
    ttest_table: pd.DataFrame,
    expert_table: pd.DataFrame,
    usage_corr_table: pd.DataFrame,
    ancova_table: pd.DataFrame,
    prepost_summary: pd.DataFrame,
    validation_summary: pd.DataFrame,
    validation_group_summary: pd.DataFrame,
    inter_rater_summary: pd.DataFrame,
) -> None:
    title, sections, results_text = build_paper_sections(
        participants,
        group_table,
        ttest_table,
        expert_table,
        usage_corr_table,
        ancova_table,
        prepost_summary,
        validation_summary,
        validation_group_summary,
        inter_rater_summary,
    )

    markdown_lines = [f"# {title}", ""]
    text_lines = [title, "=" * len(title), ""]

    for section_title, paragraphs in sections:
        markdown_lines.extend([f"## {section_title}", ""])
        text_lines.extend([section_title, "-" * len(section_title)])
        for paragraph in paragraphs:
            markdown_lines.extend([paragraph, ""])
            text_lines.extend([paragraph, ""])

    (output_dir / "paper_draft_zh.md").write_text("\n".join(markdown_lines).strip() + "\n", encoding="utf-8")
    (output_dir / "paper_draft_zh.txt").write_text("\n".join(text_lines).strip() + "\n", encoding="utf-8")
    (output_dir / "results_section_zh.txt").write_text(results_text.strip() + "\n", encoding="utf-8")

    if Document is not None:
        document = Document()
        document.add_heading(title, level=0)
        for section_title, paragraphs in sections:
            document.add_heading(section_title, level=1)
            for paragraph in paragraphs:
                document.add_paragraph(paragraph)
        document.save(output_dir / "paper_draft_zh.docx")


def write_summary_report(
    output_dir: Path,
    participants: pd.DataFrame,
    group_table: pd.DataFrame,
    ttest_table: pd.DataFrame,
    expert_table: pd.DataFrame,
    usage_corr_table: pd.DataFrame,
    ancova_table: pd.DataFrame,
    prepost_summary: pd.DataFrame,
    validation_summary: pd.DataFrame,
    validation_group_summary: pd.DataFrame,
    inter_rater_summary: pd.DataFrame,
    inter_rater_by_group: pd.DataFrame,
    inter_rater_by_stage: pd.DataFrame,
    inter_rater_by_piece: pd.DataFrame,
    adjudication_queue: pd.DataFrame,
) -> None:
    participants = ensure_columns(participants, ["groupId"])
    participant_count = len(participants)
    experimental_count = int((participants["groupId"] == "experimental").sum())
    control_count = int((participants["groupId"] == "control").sum())

    group_lookup = {row["groupId"]: row for _, row in group_table.iterrows()}
    exp_group = group_lookup.get("experimental", {})
    ctl_group = group_lookup.get("control", {})

    lines = [
        "# Research Summary",
        "",
        f"- Participants: {participant_count}",
        f"- Experimental group: {experimental_count}",
        f"- Control group: {control_count}",
        "",
        "## Group Means",
        f"- Experimental pitch gain mean: {round_value(exp_group.get('pitchGainMean'), 2)}",
        f"- Control pitch gain mean: {round_value(ctl_group.get('pitchGainMean'), 2)}",
        f"- Experimental rhythm gain mean: {round_value(exp_group.get('rhythmGainMean'), 2)}",
        f"- Control rhythm gain mean: {round_value(ctl_group.get('rhythmGainMean'), 2)}",
        f"- Experimental usefulness mean: {round_value(exp_group.get('usefulnessMean'), 2)}",
        f"- Control usefulness mean: {round_value(ctl_group.get('usefulnessMean'), 2)}",
        "",
        "## Group Tests",
    ]

    for _, row in ttest_table.iterrows():
        lines.append(
            f"- {row['metric']}: experimental={round_value(row['experimentalMean'], 3)}, "
            f"control={round_value(row['controlMean'], 3)}, "
            f"t={round_value(row['tStatistic'], 3)}, p={round_value(row['pValue'], 4)}"
        )

    lines.extend(["", "## Pre/Post Summary"])
    if prepost_summary.empty:
        lines.append("- No paired score data available yet.")
    else:
        for _, row in prepost_summary.iterrows():
            lines.append(
                f"- {row['metric']} | {row['groupId']} | {row['time']}: "
                f"n={int(row['sampleSize'])}, mean={round_value(row['meanScore'], 3)}, sem={round_value(row['sem'], 3)}"
            )

    lines.extend(["", "## ANCOVA"])
    for _, row in ancova_table.iterrows():
        lines.append(
            f"- {row['metric']} {row['term']}: "
            f"n={int(row['sampleSize']) if pd.notna(row['sampleSize']) else 0}, "
            f"F={round_value(row['fStatistic'], 3)}, p={round_value(row['pValue'], 4)}, "
            f"partial_eta_sq={round_value(row['partialEtaSquared'], 4)}, "
            f"adj_R2={round_value(row['adjustedRsquared'], 4)}"
        )

    lines.extend(["", "## System vs Expert Correlations"])
    for _, row in expert_table.iterrows():
        lines.append(
            f"- {row['systemMetric']} vs {row['expertMetric']}: "
            f"n={int(row['sampleSize'])}, r={round_value(row['pearsonR'], 3)}, p={round_value(row['pValue'], 4)}"
        )

    lines.extend(["", "## Teacher Validation Alignment"])
    if validation_summary.empty or not int(validation_summary.iloc[0].get("reviewCount", 0) or 0):
        lines.append("- No teacher validation reviews available yet.")
    else:
        row = validation_summary.iloc[0]
        lines.extend(
            [
                f"- Reviews: {int(row['reviewCount'])}",
                f"- Participants covered: {int(row['participantCount'])}",
                f"- Mean agreement (1-5): {round_value(row['averageAgreement'], 3)}",
                f"- Mean note F1: {round_value(row['averageNoteF1'], 3)}",
                f"- Mean measure F1: {round_value(row['averageMeasureF1'], 3)}",
                f"- Practice-path agreement rate: {round_value((row['pathAgreementRate'] or 0) * 100, 2)}%",
            ]
        )
        for _, group_row in validation_group_summary.iterrows():
            lines.append(
                f"- {group_row['groupId']}: reviews={int(group_row['reviewCount'])}, "
                f"note_f1={round_value(group_row['averageNoteF1'], 3)}, "
                f"measure_f1={round_value(group_row['averageMeasureF1'], 3)}, "
                f"path_agreement={round_value((group_row['pathAgreementRate'] or 0) * 100, 2)}%"
            )

    lines.extend(["", "## Inter-rater Reliability"])
    if inter_rater_summary.empty or not int(inter_rater_summary.iloc[0].get("pairCount", 0) or 0):
        lines.append("- No dual-rater teacher reviews available yet.")
    else:
        row = inter_rater_summary.iloc[0]
        lines.extend(
            [
                f"- Dual-rated pairs: {int(row['pairCount'])}",
                f"- Analyses covered: {int(row['analysisCount'])}",
                f"- Participants covered: {int(row['participantCount'])}",
                f"- Cohen's kappa (practice path): {round_value(row['pathCohenKappa'], 3)}",
                f"- ICC (overall agreement): {round_value(row['overallAgreementICC'], 3)}",
                f"- Practice-path match rate: {round_value((row['pathMatchRate'] or 0) * 100, 2)}%",
                f"- Mean agreement gap: {round_value(row['meanAgreementGap'], 3)}",
                f"- Mean note overlap F1: {round_value(row['meanNoteOverlapF1'], 3)}",
                f"- Mean measure overlap F1: {round_value(row['meanMeasureOverlapF1'], 3)}",
                f"- Adjudication queue: {int(row['adjudicationCount'])} pair(s), rate={round_value((row['adjudicationRate'] or 0) * 100, 2)}%",
            ]
        )
    if not inter_rater_by_group.empty:
        lines.extend(["", "## Inter-rater Reliability by Group"])
        for _, row in inter_rater_by_group.iterrows():
            lines.append(
                f"- {row['groupId']}: pairs={int(row['pairCount'])}, "
                f"kappa={round_value(row['pathCohenKappa'], 3)}, "
                f"icc={round_value(row['overallAgreementICC'], 3)}, "
                f"path_match={round_value((row['pathMatchRate'] or 0) * 100, 2)}%, "
                f"adjudication_rate={round_value((row['adjudicationRate'] or 0) * 100, 2)}%"
            )
    if not inter_rater_by_stage.empty:
        lines.extend(["", "## Inter-rater Reliability by Stage"])
        for _, row in inter_rater_by_stage.iterrows():
            lines.append(
                f"- {row['sessionStage']}: pairs={int(row['pairCount'])}, "
                f"kappa={round_value(row['pathCohenKappa'], 3)}, "
                f"icc={round_value(row['overallAgreementICC'], 3)}, "
                f"note_f1={round_value(row['meanNoteOverlapF1'], 3)}, "
                f"adjudication_rate={round_value((row['adjudicationRate'] or 0) * 100, 2)}%"
            )
    if not inter_rater_by_piece.empty:
        lines.extend(["", "## Inter-rater Reliability by Piece"])
        for _, row in inter_rater_by_piece.head(8).iterrows():
            lines.append(
                f"- {row['scoreUnit']}: pairs={int(row['pairCount'])}, "
                f"kappa={round_value(row['pathCohenKappa'], 3)}, "
                f"icc={round_value(row['overallAgreementICC'], 3)}, "
                f"note_f1={round_value(row['meanNoteOverlapF1'], 3)}, "
                f"adjudication_rate={round_value((row['adjudicationRate'] or 0) * 100, 2)}%"
            )
    lines.extend(["", "## Adjudication Queue"])
    if adjudication_queue.empty:
        lines.append("- No pairs currently meet the adjudication trigger rules.")
    else:
        for _, row in adjudication_queue.iterrows():
            lines.append(
                f"- {row['analysisId']} | {row['scoreUnit']} | {row['raterAId']} vs {row['raterBId']} | "
                f"reasons={row['adjudicationReason']}"
            )

    lines.extend(["", "## Usage Correlations"])
    for _, row in usage_corr_table.iterrows():
        lines.append(
            f"- analysisCount vs {row['outcome']}: "
            f"n={int(row['sampleSize'])}, r={round_value(row['pearsonR'], 3)}, p={round_value(row['pValue'], 4)}"
        )

    (output_dir / "report.md").write_text("\n".join(lines), encoding="utf-8")


def write_summary_json(
    output_dir: Path,
    group_table: pd.DataFrame,
    ttest_table: pd.DataFrame,
    expert_table: pd.DataFrame,
    usage_corr_table: pd.DataFrame,
    ancova_table: pd.DataFrame,
    prepost_summary: pd.DataFrame,
    validation_summary: pd.DataFrame,
    validation_group_summary: pd.DataFrame,
    validation_path_confusion: pd.DataFrame,
    inter_rater_pairs: pd.DataFrame,
    inter_rater_summary: pd.DataFrame,
    inter_rater_by_group: pd.DataFrame,
    inter_rater_by_stage: pd.DataFrame,
    inter_rater_by_piece: pd.DataFrame,
    adjudication_queue: pd.DataFrame,
) -> None:
    payload = {
        "groupSummary": group_table.to_dict(orient="records"),
        "groupTests": ttest_table.to_dict(orient="records"),
        "prePostSummary": prepost_summary.to_dict(orient="records"),
        "ancova": ancova_table.to_dict(orient="records"),
        "systemExpertCorrelations": expert_table.to_dict(orient="records"),
        "usageCorrelations": usage_corr_table.to_dict(orient="records"),
        "validationSummary": validation_summary.to_dict(orient="records"),
        "validationByGroup": validation_group_summary.to_dict(orient="records"),
        "validationPathConfusion": validation_path_confusion.to_dict(orient="records"),
        "interRaterPairs": inter_rater_pairs.to_dict(orient="records"),
        "interRaterSummary": inter_rater_summary.to_dict(orient="records"),
        "interRaterByGroup": inter_rater_by_group.to_dict(orient="records"),
        "interRaterByStage": inter_rater_by_stage.to_dict(orient="records"),
        "interRaterByPiece": inter_rater_by_piece.to_dict(orient="records"),
        "adjudicationQueue": adjudication_queue.to_dict(orient="records"),
    }
    (output_dir / "summary.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def plot_gain_by_group(participants: pd.DataFrame, output_dir: Path) -> None:
    participants = ensure_columns(participants, ["participantId", "groupId", "pitchGain", "rhythmGain"])
    plot_frame = participants.melt(
        id_vars=["participantId", "groupId"],
        value_vars=["pitchGain", "rhythmGain"],
        var_name="metric",
        value_name="value",
    ).dropna()
    if plot_frame.empty:
        save_figure(placeholder_figure("System Gains by Group"), output_dir, "figure_gain_by_group")
        return

    fig, ax = plt.subplots(figsize=(9, 5))
    sns.boxplot(data=plot_frame, x="metric", y="value", hue="groupId", ax=ax)
    sns.stripplot(data=plot_frame, x="metric", y="value", hue="groupId", dodge=True, ax=ax, alpha=0.45, linewidth=0)
    handles, labels = ax.get_legend_handles_labels()
    ax.legend(handles[:2], labels[:2], title="group")
    ax.set_title("System Gains by Group")
    ax.set_xlabel("")
    ax.set_ylabel("Gain")
    save_figure(fig, output_dir, "figure_gain_by_group")


def plot_questionnaire_bars(questionnaires: pd.DataFrame, output_dir: Path) -> None:
    questionnaires = ensure_columns(questionnaires, ["groupId", *QUESTIONNAIRE_NUMERIC_COLUMNS])
    summary = questionnaires.groupby("groupId", dropna=False)[QUESTIONNAIRE_NUMERIC_COLUMNS].mean().reset_index()
    plot_frame = summary.melt(id_vars=["groupId"], value_vars=QUESTIONNAIRE_NUMERIC_COLUMNS, var_name="metric", value_name="score").dropna()
    if plot_frame.empty:
        save_figure(placeholder_figure("Questionnaire Means by Group"), output_dir, "figure_questionnaire_by_group")
        return

    fig, ax = plt.subplots(figsize=(10, 5))
    sns.barplot(data=plot_frame, x="metric", y="score", hue="groupId", ax=ax)
    ax.set_title("Questionnaire Means by Group")
    ax.set_xlabel("")
    ax.set_ylabel("Mean score")
    ax.set_ylim(0, 5.2)
    save_figure(fig, output_dir, "figure_questionnaire_by_group")


def plot_system_vs_expert(participants: pd.DataFrame, output_dir: Path) -> None:
    participants = ensure_columns(participants, ["posttestPitch", "expertPosttestPitch", "groupId"])
    subset = participants[["posttestPitch", "expertPosttestPitch", "groupId"]].dropna()
    if subset.empty:
        save_figure(placeholder_figure("System vs Expert Posttest Pitch"), output_dir, "figure_system_vs_expert")
        return

    fig, ax = plt.subplots(figsize=(6, 6))
    sns.scatterplot(data=subset, x="posttestPitch", y="expertPosttestPitch", hue="groupId", ax=ax, s=70)
    if len(subset) >= 2 and subset["posttestPitch"].nunique() > 1 and subset["expertPosttestPitch"].nunique() > 1:
        sns.regplot(data=subset, x="posttestPitch", y="expertPosttestPitch", scatter=False, ax=ax, color="#444444")
    ax.set_title("System vs Expert Posttest Pitch")
    ax.set_xlabel("System score")
    ax.set_ylabel("Expert score")
    save_figure(fig, output_dir, "figure_system_vs_expert")


def plot_usage_vs_gain(participants: pd.DataFrame, output_dir: Path) -> None:
    participants = ensure_columns(participants, ["analysisCount", "pitchGain", "groupId"])
    subset = participants[["analysisCount", "pitchGain", "groupId"]].dropna()
    if subset.empty:
        save_figure(placeholder_figure("Analysis Usage vs Pitch Gain"), output_dir, "figure_usage_vs_pitch_gain")
        return

    fig, ax = plt.subplots(figsize=(7, 5))
    sns.scatterplot(data=subset, x="analysisCount", y="pitchGain", hue="groupId", ax=ax, s=70)
    if len(subset) >= 2 and subset["analysisCount"].nunique() > 1 and subset["pitchGain"].nunique() > 1:
        sns.regplot(data=subset, x="analysisCount", y="pitchGain", scatter=False, ax=ax, color="#444444")
    ax.set_title("Analysis Usage vs Pitch Gain")
    ax.set_xlabel("Analysis count")
    ax.set_ylabel("Pitch gain")
    save_figure(fig, output_dir, "figure_usage_vs_pitch_gain")


def plot_prepost_trends(prepost_summary: pd.DataFrame, output_dir: Path) -> None:
    prepost_summary = ensure_columns(prepost_summary, ["metric", "groupId", "time", "meanScore", "sem"])
    if prepost_summary.empty:
        save_figure(placeholder_figure("Pre/Post Trends by Group"), output_dir, "figure_prepost_trends")
        return

    fig, axes = plt.subplots(1, 2, figsize=(11, 4.5), sharey=True)
    metric_titles = {"pitch": "Pitch score", "rhythm": "Rhythm score"}
    x_lookup = {label: index for index, label in enumerate(PREPOST_TIME_ORDER)}

    for ax, metric in zip(axes, ["pitch", "rhythm"]):
        metric_data = prepost_summary.loc[prepost_summary["metric"] == metric].copy()
        if metric_data.empty:
            add_placeholder_axis(ax, metric_titles[metric])
            continue

        for group_id, group_data in metric_data.groupby("groupId", dropna=False):
            group_data = group_data.sort_values("time")
            xs = [x_lookup[str(item)] for item in group_data["time"]]
            ys = group_data["meanScore"].to_list()
            errs = [0 if pd.isna(value) else value for value in group_data["sem"]]
            ax.errorbar(xs, ys, yerr=errs, marker="o", linewidth=2, capsize=4, label=str(group_id))

        ax.set_title(metric_titles[metric])
        ax.set_xticks(list(x_lookup.values()), PREPOST_TIME_ORDER)
        ax.set_xlabel("")
        ax.set_ylabel("Mean score")
        ax.set_ylim(0, 100)
        ax.legend(title="group")

    save_figure(fig, output_dir, "figure_prepost_trends")


def plot_validation_metrics(validation_group_summary: pd.DataFrame, output_dir: Path) -> None:
    validation_group_summary = ensure_columns(
        validation_group_summary,
        ["groupId", "averageNoteF1", "averageMeasureF1", "pathAgreementRate", "averageAgreementNormalized"],
    )
    plot_frame = validation_group_summary.melt(
        id_vars=["groupId"],
        value_vars=["averageNoteF1", "averageMeasureF1", "pathAgreementRate", "averageAgreementNormalized"],
        var_name="metric",
        value_name="value",
    ).dropna()
    if plot_frame.empty:
        save_figure(placeholder_figure("Teacher Validation Metrics by Group"), output_dir, "figure_validation_by_group")
        return

    fig, ax = plt.subplots(figsize=(9, 5))
    sns.barplot(data=plot_frame, x="metric", y="value", hue="groupId", ax=ax)
    ax.set_title("Teacher Validation Metrics by Group")
    ax.set_xlabel("")
    ax.set_ylabel("Mean proportion")
    ax.set_ylim(0, 1.05)
    save_figure(fig, output_dir, "figure_validation_by_group")


def plot_validation_path_heatmap(validation_path_confusion: pd.DataFrame, output_dir: Path) -> None:
    validation_path_confusion = ensure_columns(validation_path_confusion, ["teacherPrimaryPath", "systemRecommendedPath", "reviewCount"])
    if validation_path_confusion.empty:
        save_figure(placeholder_figure("Teacher vs System Practice Path"), output_dir, "figure_validation_path_heatmap")
        return

    pivot = (
        validation_path_confusion.pivot(index="teacherPrimaryPath", columns="systemRecommendedPath", values="reviewCount")
        .reindex(index=PRACTICE_PATH_ORDER, columns=PRACTICE_PATH_ORDER)
        .fillna(0)
    )
    fig, ax = plt.subplots(figsize=(6.5, 5.2))
    sns.heatmap(pivot, annot=True, fmt=".0f", cmap="Blues", cbar=True, ax=ax)
    ax.set_title("Teacher vs System Practice Path")
    ax.set_xlabel("System recommended path")
    ax.set_ylabel("Teacher primary path")
    save_figure(fig, output_dir, "figure_validation_path_heatmap")


def plot_inter_rater_metrics(inter_rater_summary: pd.DataFrame, output_dir: Path) -> None:
    inter_rater_summary = ensure_columns(
        inter_rater_summary,
        ["pathCohenKappa", "overallAgreementICC", "meanNoteOverlapF1", "meanMeasureOverlapF1"],
    )
    if inter_rater_summary.empty or not inter_rater_summary.iloc[0].notna().any():
        save_figure(placeholder_figure("Inter-rater Reliability Metrics"), output_dir, "figure_inter_rater_metrics")
        return

    row = inter_rater_summary.iloc[0]
    plot_frame = pd.DataFrame(
        {
            "metric": ["pathCohenKappa", "overallAgreementICC", "meanNoteOverlapF1", "meanMeasureOverlapF1"],
            "value": [row["pathCohenKappa"], row["overallAgreementICC"], row["meanNoteOverlapF1"], row["meanMeasureOverlapF1"]],
        }
    ).dropna()
    if plot_frame.empty:
        save_figure(placeholder_figure("Inter-rater Reliability Metrics"), output_dir, "figure_inter_rater_metrics")
        return

    fig, ax = plt.subplots(figsize=(8, 4.8))
    sns.barplot(data=plot_frame, x="metric", y="value", ax=ax, color="#4f46e5")
    ax.set_title("Inter-rater Reliability Metrics")
    ax.set_xlabel("")
    ax.set_ylabel("Coefficient")
    ax.set_ylim(0, 1.05)
    save_figure(fig, output_dir, "figure_inter_rater_metrics")


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate SSCI-ready tables and figures from exported study CSV files.")
    parser.add_argument("--participants", required=True, type=Path)
    parser.add_argument("--questionnaires", required=True, type=Path)
    parser.add_argument("--ratings", required=True, type=Path)
    parser.add_argument("--analyses", required=True, type=Path)
    parser.add_argument("--validations", required=False, type=Path, default=None)
    parser.add_argument("--output-dir", default=Path("research-analysis/output"), type=Path)
    args = parser.parse_args()

    output_dir = args.output_dir
    output_dir.mkdir(parents=True, exist_ok=True)
    sns.set_theme(style="whitegrid", palette="deep")

    participants = ensure_numeric(safe_read_csv(args.participants), PARTICIPANT_NUMERIC_COLUMNS)
    questionnaires = ensure_numeric(safe_read_csv(args.questionnaires), QUESTIONNAIRE_NUMERIC_COLUMNS)
    ratings = ensure_numeric(safe_read_csv(args.ratings), RATING_NUMERIC_COLUMNS)
    analyses = ensure_numeric(safe_read_csv(args.analyses), ANALYSIS_NUMERIC_COLUMNS)
    validation_reviews = prepare_validation_reviews(
        safe_read_optional_csv(
            args.validations,
            columns=[
                "reviewId",
                "analysisId",
                "participantId",
                "groupId",
                "sessionStage",
                "pieceId",
                "sectionId",
                "raterId",
                "overallAgreement",
                "teacherPrimaryPath",
                "systemRecommendedPath",
                "pathAgreement",
                "noteMatchedCount",
                "notePrecision",
                "noteRecall",
                "noteF1",
                "measureMatchedCount",
                "measurePrecision",
                "measureRecall",
                "measureF1",
                "submittedAt",
            ],
        )
    )

    group_table = group_summary(participants)
    ttest_table = ttest_summary(participants)
    expert_table = build_expert_system_table(participants)
    usage_table = build_analysis_usage_table(analyses)
    usage_corr_table = build_usage_correlation_table(participants)
    prepost_long = build_prepost_long_table(participants)
    prepost_summary = build_prepost_summary(prepost_long)
    ancova_table = build_ancova_table(participants)
    validation_summary = build_validation_summary(validation_reviews)
    validation_group_summary = build_validation_group_summary(validation_reviews)
    validation_path_confusion = build_validation_path_confusion(validation_reviews)
    inter_rater_pairs = build_inter_rater_pairs(validation_reviews)
    inter_rater_summary = build_inter_rater_summary(inter_rater_pairs)
    inter_rater_by_group = build_inter_rater_breakdown(inter_rater_pairs, ["groupId"])
    inter_rater_by_stage = build_inter_rater_breakdown(inter_rater_pairs, ["sessionStage"])
    inter_rater_by_piece = build_inter_rater_breakdown(inter_rater_pairs, ["scoreUnit"])
    adjudication_queue = build_inter_rater_adjudication_queue(inter_rater_pairs)

    save_table(participants, output_dir, "table_participant_overview")
    save_table(group_table, output_dir, "table_group_summary")
    save_table(ttest_table, output_dir, "table_group_ttests")
    save_table(questionnaire_summary(questionnaires), output_dir, "table_questionnaire_summary")
    save_table(expert_table, output_dir, "table_system_expert_correlations")
    save_table(usage_table, output_dir, "table_analysis_usage")
    save_table(usage_corr_table, output_dir, "table_usage_correlations")
    save_table(prepost_long, output_dir, "table_prepost_long")
    save_table(prepost_summary, output_dir, "table_prepost_summary")
    save_table(ancova_table, output_dir, "table_ancova_summary")
    save_table(ratings.sort_values(["participantId", "stage", "submittedAt"]), output_dir, "table_expert_ratings_raw")
    save_table(validation_reviews.sort_values(["participantId", "submittedAt"]), output_dir, "table_validation_reviews_raw")
    save_table(validation_summary, output_dir, "table_validation_summary")
    save_table(validation_group_summary, output_dir, "table_validation_group_summary")
    save_table(validation_path_confusion, output_dir, "table_validation_path_confusion")
    save_table(inter_rater_pairs, output_dir, "table_inter_rater_pairs")
    save_table(inter_rater_summary, output_dir, "table_inter_rater_summary")
    save_table(inter_rater_by_group, output_dir, "table_inter_rater_by_group")
    save_table(inter_rater_by_stage, output_dir, "table_inter_rater_by_stage")
    save_table(inter_rater_by_piece, output_dir, "table_inter_rater_by_piece")
    save_table(adjudication_queue, output_dir, "table_inter_rater_adjudication_queue")

    plot_gain_by_group(participants, output_dir)
    plot_questionnaire_bars(questionnaires, output_dir)
    plot_system_vs_expert(participants, output_dir)
    plot_usage_vs_gain(participants, output_dir)
    plot_prepost_trends(prepost_summary, output_dir)
    plot_validation_metrics(validation_group_summary, output_dir)
    plot_validation_path_heatmap(validation_path_confusion, output_dir)
    plot_inter_rater_metrics(inter_rater_summary, output_dir)
    write_summary_report(
        output_dir,
        participants,
        group_table,
        ttest_table,
        expert_table,
        usage_corr_table,
        ancova_table,
        prepost_summary,
        validation_summary,
        validation_group_summary,
        inter_rater_summary,
        inter_rater_by_group,
        inter_rater_by_stage,
        inter_rater_by_piece,
        adjudication_queue,
    )
    write_summary_json(
        output_dir,
        group_table,
        ttest_table,
        expert_table,
        usage_corr_table,
        ancova_table,
        prepost_summary,
        validation_summary,
        validation_group_summary,
        validation_path_confusion,
        inter_rater_pairs,
        inter_rater_summary,
        inter_rater_by_group,
        inter_rater_by_stage,
        inter_rater_by_piece,
        adjudication_queue,
    )
    write_paper_draft(
        output_dir,
        participants,
        group_table,
        ttest_table,
        expert_table,
        usage_corr_table,
        ancova_table,
        prepost_summary,
        validation_summary,
        validation_group_summary,
        inter_rater_summary,
    )

    summary_path = output_dir / "README.txt"
    summary_path.write_text(
        "\n".join(
            [
                "Generated files:",
                "- table_participant_overview.csv",
                "- table_group_summary.csv",
                "- table_group_ttests.csv",
                "- table_questionnaire_summary.csv",
                "- table_system_expert_correlations.csv",
                "- table_analysis_usage.csv",
                "- table_usage_correlations.csv",
                "- table_prepost_long.csv",
                "- table_prepost_summary.csv",
                "- table_ancova_summary.csv",
                "- table_expert_ratings_raw.csv",
                "- table_validation_reviews_raw.csv",
                "- table_validation_summary.csv",
                "- table_validation_group_summary.csv",
                "- table_validation_path_confusion.csv",
                "- table_inter_rater_pairs.csv",
                "- table_inter_rater_summary.csv",
                "- table_inter_rater_by_group.csv",
                "- table_inter_rater_by_stage.csv",
                "- table_inter_rater_by_piece.csv",
                "- table_inter_rater_adjudication_queue.csv",
                "- figure_gain_by_group.png",
                "- figure_questionnaire_by_group.png",
                "- figure_system_vs_expert.png",
                "- figure_usage_vs_pitch_gain.png",
                "- figure_prepost_trends.png",
                "- figure_validation_by_group.png",
                "- figure_validation_path_heatmap.png",
                "- figure_inter_rater_metrics.png",
                "- report.md",
                "- summary.json",
                "- paper_draft_zh.md",
                "- paper_draft_zh.txt",
                "- paper_draft_zh.docx",
                "- results_section_zh.txt",
            ]
        ),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
